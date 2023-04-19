import datetime
from docx.shared import Inches
from naming import Heb
from order import Order
from config_word import setup_doc


def make_orders_doc(orders: list[Order]):
    document, rtl2_style = setup_doc()

    for order in orders:
        # creating word page for each order
        order_doc = order.to_dict()

        # title
        document.add_paragraph(order_doc.get('title'), style='rtl')

        # address
        if order_doc.get('is_delivery'):
            address = order_doc.get('address', '')
            document.add_paragraph(address, style='rtl')

        document.add_paragraph(order_doc.get('comment', ''), style='rtl_red')

        produces = order_doc.get('produces')
        # adding the products table
        table = document.add_table(rows=1, cols=5, style='Light Grid Accent 1')
        table.allow_autofit = False
        hdr_cells = table.rows[0].cells
        hdr_cells[4].text = Heb.product_name
        hdr_cells[3].text = Heb.package
        hdr_cells[2].text = Heb.unit_price
        hdr_cells[1].text = Heb.amount
        hdr_cells[0].text = Heb.total
        # adding style to the table header
        for i in range(5):
            hdr_cells[i].paragraphs[0].paragraph_format.alignment = 2
            hdr_cells[i].paragraphs[0].style = rtl2_style

        for product_name, package, unit_price, amount, product_sum in produces:
            row_cells = table.add_row().cells
            row_cells[0].text = str(product_sum)
            row_cells[1].text = str(amount)
            row_cells[2].text = str(unit_price)
            row_cells[3].text = package
            row_cells[4].text = product_name
            # adding style to the table row
            for i in range(5):
                row_cells[i].paragraphs[0].paragraph_format.alignment = 2
                row_cells[i].paragraphs[0].style = rtl2_style

        widths = (Inches(1), Inches(1), Inches(2), Inches(3), Inches(5))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        # adding page break
        document.add_page_break()
    today = str(datetime.date.today())
    result_file_name = f'{Heb.output_name}_{today}.docx'
    document.save(result_file_name)
    print(f"orders doc is in : {result_file_name}")